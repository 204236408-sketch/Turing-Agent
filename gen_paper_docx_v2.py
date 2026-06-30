import os
import re
import base64
from io import BytesIO
from mistune import create_markdown, HTMLRenderer
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PAPER_DIR = os.path.join(os.path.dirname(__file__), "paper")
CHAPTER_FILES = [
    "ch00_abstract.md",
    "ch01_intro.md",
    "ch02_requirements.md",
    "ch03_architecture.md",
    "ch04_database.md",
    "ch05_agent.md",
    "ch06_core_services.md",
    "ch07_prompt_engineering.md",
    "ch08_implementation.md",
    "ch09_testing.md",
    "ch10_conclusion.md",
]


def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


class DocxRenderer:
    def __init__(self, doc, images_dir):
        self.doc = doc
        self.images_dir = images_dir
        self._in_table = False
        self._table_headers = []
        self._table_rows = []
        self._code_buffer = []

    def add_heading(self, text, level):
        if level == 1:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(20)
            pf.space_after = Pt(14)
            pf.line_spacing = 1.5
            run = p.add_run(text)
            set_run_font(run, cn_font="黑体", size=18, bold=True)
        elif level == 2:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(16)
            pf.space_after = Pt(8)
            pf.line_spacing = 1.5
            run = p.add_run(text)
            set_run_font(run, cn_font="黑体", size=15, bold=True)
        elif level == 3:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.5
            run = p.add_run(text)
            set_run_font(run, cn_font="黑体", size=13, bold=True)
        else:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(8)
            pf.space_after = Pt(4)
            pf.line_spacing = 1.5
            run = p.add_run(text)
            set_run_font(run, cn_font="黑体", size=12, bold=True)

    def add_paragraph(self, text, bold=False, first_indent=True, font_size=12, alignment=None, cn_font="宋体"):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        if first_indent:
            pf.first_line_indent = Cm(0.74)
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.5
        if alignment:
            p.alignment = alignment
        run = p.add_run(text)
        set_run_font(run, cn_font=cn_font, size=font_size, bold=bold)
        return p

    def render_text(self, text):
        """Simple text paragraph."""
        self.add_paragraph(text)

    def render_bold_text(self, text):
        self.add_paragraph(text, bold=True)

    def render_table(self, headers, rows):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h.strip())
            set_run_font(run, size=9, bold=True)
            set_cell_shading(cell, "D9E2F3")

        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < len(headers):
                    cell = table.rows[r_idx + 1].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(cell_text.strip())
                    set_run_font(run, size=9)
                    if r_idx % 2 == 1:
                        set_cell_shading(cell, "F2F2F2")

        self.add_paragraph("", font_size=2, first_indent=False)

    def render_image(self, src, alt="", width=5.0):
        img_path = os.path.join(self.images_dir, src)
        if os.path.exists(img_path):
            try:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(width))
            except Exception:
                self.add_paragraph(f"[图片: {alt}]", first_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)
        else:
            self.add_paragraph(f"[图片: {alt} ({src})]", first_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)

    def render_code_block(self, code_text):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.0
        pf.left_indent = Cm(0.5)
        run = p.add_run(code_text)
        set_run_font(run, cn_font="宋体", en_font="Courier New", size=8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
        pPr = p._element.get_or_add_pPr()
        pPr.append(shading)

    def render_html_caption(self, text):
        clean = re.sub(r'<[^>]+>', '', text).strip()
        if clean:
            self.add_paragraph(clean, first_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10)


def parse_markdown_tables(lines, start):
    """Parse markdown table from lines starting at start. Returns (headers, rows, next_line)."""
    if not lines[start].strip().startswith("|"):
        return None, None, start
    header_line = lines[start].strip()
    headers = [h.strip() for h in header_line.split("|")[1:-1]]
    i = start + 1
    # skip separator line
    if i < len(lines) and "---" in lines[i]:
        i += 1
    rows = []
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return headers, rows, i


def generate_docx():
    doc = Document()

    # Page setup (A4, standard margins)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    for fname in CHAPTER_FILES:
        fpath = os.path.join(PAPER_DIR, fname)
        if not os.path.exists(fpath):
            continue
        with open(fpath, "r", encoding="utf-8") as f:
            text = f.read()

        lines = text.split("\n")
        i = 0
        in_code = False
        code_lines = []

        while i < len(lines):
            raw = lines[i]
            stripped = raw.strip()

            # Code block
            if stripped.startswith("```"):
                if not in_code:
                    in_code = True
                    code_lines = []
                    i += 1
                    continue
                else:
                    in_code = False
                    code_text = "\n".join(code_lines)
                    renderer = DocxRenderer(doc, PAPER_DIR)
                    renderer.render_code_block(code_text)
                    i += 1
                    continue

            if in_code:
                code_lines.append(raw)
                i += 1
                continue

            # HTML image tag
            html_img = re.search(r'<img\s+src="([^"]+)"', stripped)
            if html_img:
                renderer = DocxRenderer(doc, PAPER_DIR)
                renderer.render_image(html_img.group(1))
                i += 1
                continue

            # HTML caption
            if stripped.startswith("<p") and ("small" in stripped or "center" in stripped or "align=" in stripped):
                clean = re.sub(r'<[^>]+>', '', stripped).strip()
                if clean:
                    renderer = DocxRenderer(doc, PAPER_DIR)
                    renderer.render_html_caption(stripped)
                i += 1
                continue

            # Image: ![alt](path)
            img_match = re.match(r'!\[(.*?)\]\((.+?)\)', stripped)
            if img_match:
                renderer = DocxRenderer(doc, PAPER_DIR)
                renderer.render_image(img_match.group(2), img_match.group(1))
                i += 1
                continue

            # Markdown table
            if stripped.startswith("|") and not stripped.startswith("|---"):
                headers, rows, next_i = parse_markdown_tables(lines, i)
                if headers:
                    renderer = DocxRenderer(doc, PAPER_DIR)
                    renderer.render_table(headers, rows)
                    i = next_i
                    continue

            # Empty line
            if not stripped:
                i += 1
                continue

            # Heading
            heading = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if heading:
                level = len(heading.group(1))
                text_content = heading.group(2)
                renderer = DocxRenderer(doc, PAPER_DIR)
                renderer.add_heading(text_content, level)
                i += 1
                continue

            # Bold line
            if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
                bold_text = stripped.strip("*")
                renderer = DocxRenderer(doc, PAPER_DIR)
                renderer.add_paragraph(bold_text, bold=True)
                i += 1
                continue

            # Normal paragraph
            renderer = DocxRenderer(doc, PAPER_DIR)
            renderer.add_paragraph(stripped)
            i += 1

    out_path = os.path.join(os.path.dirname(__file__), "论文-重生之我是图灵.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}  ({os.path.getsize(out_path)} bytes)")


if __name__ == "__main__":
    generate_docx()
