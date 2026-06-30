import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PAPER_DIR = os.path.join(os.path.dirname(__file__), "paper")
CHAPTER_FILES = [
    "ch00_abstract.md",
    "ch01_intro.md",
    "ch02_requirements.md",
    "ch03_architecture.md",
    "ch04_database.md",
    "ch05_agent.md",
    "ch06_core_services.md",
    "ch07_prompt_engineering.md",
    "ch08_implementation.md",
    "ch09_testing.md",
    "ch10_conclusion.md",
]


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=12, bold=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)


def add_formatted_paragraph(doc, text, style_name="Normal", alignment=None,
                            space_before=0, space_after=6, first_line_indent=None,
                            font_size=12, bold=False, cn_font="宋体", en_font="Times New Roman"):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, font_size, bold)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if alignment is not None:
        pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    pf.line_spacing = 1.5
    return p


def parse_table_from_md(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (table_data, end_idx)."""
    header_line = lines[start_idx].strip()
    headers = [h.strip() for h in header_line.split("|")[1:-1]]
    i = start_idx + 1
    rows = []
    while i < len(lines):
        line = lines[i].strip()
        if not line or not line.startswith("|"):
            break
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return headers, rows, i


def create_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── Style definitions ──
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ── Parse and add content ──
    for ch_idx, fname in enumerate(CHAPTER_FILES):
        fpath = os.path.join(PAPER_DIR, fname)
        if not os.path.exists(fpath):
            continue
        with open(fpath, "r", encoding="utf-8") as f:
            text = f.read()

        lines = text.split("\n")
        i = 0
        in_code_block = False
        code_buffer = []

        while i < len(lines):
            raw = lines[i]
            stripped = raw.strip()

            # ── Code block ──
            if stripped.startswith("```"):
                if not in_code_block:
                    in_code_block = True
                    code_buffer = []
                    i += 1
                    continue
                else:
                    in_code_block = False
                    code_text = "\n".join(code_buffer)
                    p = add_formatted_paragraph(
                        doc, code_text, font_size=9, cn_font="宋体", en_font="Courier New",
                        space_before=3, space_after=3
                    )
                    # Add shading to code block
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    i += 1
                    continue

            if in_code_block:
                code_buffer.append(raw)
                i += 1
                continue

            # ── Image ──
            img_match = re.match(r'!\[(.*?)\]\((.+?)\)', stripped)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                full_img_path = os.path.join(PAPER_DIR, img_path)
                if os.path.exists(full_img_path):
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(full_img_path, width=Inches(5.0))
                    except Exception:
                        add_formatted_paragraph(doc, f"[图片: {alt_text}]", font_size=10)
                else:
                    add_formatted_paragraph(doc, f"[图片: {alt_text} ({img_path})]", font_size=10)
                i += 1
                continue

            # ── HTML image ──
            html_img = re.search(r'<img\s+src="(.+?)"', stripped)
            if html_img:
                img_path = html_img.group(1)
                full_img_path = os.path.join(PAPER_DIR, img_path)
                if os.path.exists(full_img_path):
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(full_img_path, width=Inches(5.0))
                    except Exception:
                        pass
                i += 1
                continue

            # ── HTML center / small text (captions) ──
            if stripped.startswith("<p align=") or stripped.startswith("<p><small>"):
                html_clean = re.sub(r'<[^>]+>', '', stripped)
                if html_clean.strip():
                    p = add_formatted_paragraph(
                        doc, html_clean.strip(), font_size=10,
                        cn_font="宋体", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=2, space_after=6
                    )
                i += 1
                continue

            # ── Markdown table ──
            if stripped.startswith("|") and not stripped.startswith("|---"):
                headers, rows, next_i = parse_table_from_md(lines, i)
                if headers:
                    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = "Table Grid"

                    # Header row
                    for j, h in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(h)
                        set_run_font(run, size=10, bold=True)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_cell_shading(cell, "D9E2F3")

                    # Data rows
                    for r_idx, row_data in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < len(headers):
                                cell = table.rows[r_idx + 1].cells[c_idx]
                                cell.text = ""
                                run = cell.paragraphs[0].add_run(cell_text)
                                set_run_font(run, size=10)
                                if r_idx % 2 == 1:
                                    set_cell_shading(cell, "F2F2F2")

                    i = next_i
                    # add spacing after table
                    add_formatted_paragraph(doc, "", font_size=6, space_before=0, space_after=0)
                    continue

            # ── Empty line ──
            if not stripped:
                i += 1
                continue

            # ── Heading detection ──
            heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                title_text = heading_match.group(2)
                if level == 1:
                    add_formatted_paragraph(
                        doc, title_text, font_size=18, bold=True, cn_font="黑体",
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=18, space_after=12
                    )
                elif level == 2:
                    add_formatted_paragraph(
                        doc, title_text, font_size=15, bold=True, cn_font="黑体",
                        space_before=14, space_after=8
                    )
                elif level == 3:
                    add_formatted_paragraph(
                        doc, title_text, font_size=13, bold=True, cn_font="黑体",
                        space_before=10, space_after=6
                    )
                elif level == 4:
                    add_formatted_paragraph(
                        doc, title_text, font_size=12, bold=True, cn_font="黑体",
                        space_before=8, space_after=4
                    )
                i += 1
                continue

            # ── Bold text (e.g., **text**) ──
            if stripped.startswith("**") and stripped.endswith("**"):
                bold_text = stripped.strip("*")
                add_formatted_paragraph(
                    doc, bold_text, font_size=12, bold=True,
                    space_before=4, space_after=4
                )
                i += 1
                continue

            # ── Normal paragraph ──
            p = add_formatted_paragraph(
                doc, stripped, font_size=12,
                first_line_indent=0.74,
                space_before=0, space_after=4
            )
            i += 1

    # ── Save ──
    out_path = os.path.join(os.path.dirname(__file__), "论文-重生之我是图灵.docx")
    doc.save(out_path)
    print(f"Document saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    create_document()
